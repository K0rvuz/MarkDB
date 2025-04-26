import sqlite3
import base64
import textwrap
import asyncio
import os
import pytesseract
import pandas as pd
import fitz  # pymupdf
import docx
import openpyxl
from pathlib import Path
from datetime import datetime
from PIL import Image
from io import BytesIO
from typing import Optional, Tuple, List, Dict, Any
from openai import OpenAI
import aiosqlite
import streamlit as st
from concurrent.futures import ThreadPoolExecutor
from tqdm import tqdm
import hashlib

# ==================== CONFIGURAÇÕES ====================
DB_PATH = "markdb.sqlite"
CHUNK_SIZE = 800  # Valor padrão, pode ser sobrescrito pelo session_state

# ==================== PROCESSAMENTO DE PDF ====================
class PDFProcessor:
    @staticmethod
    def extract_text(pdf_path: str) -> str:
        """Extrai texto de PDF com pymupdf (rápido e mantém estrutura)"""
        try:
            with fitz.open(pdf_path) as doc:
                return "\n".join(page.get_text() for page in doc)
        except Exception as e:
            print(f"Erro ao processar {pdf_path}: {str(e)}")
            return ""

    @staticmethod
    def get_cache_path(pdf_path: str, cache_dir: str = ".markdb_cache") -> Path:
        """Gera caminho de cache usando hash do arquivo"""
        Path(cache_dir).mkdir(exist_ok=True)
        file_hash = hashlib.md5(open(pdf_path, 'rb').read()).hexdigest()
        return Path(cache_dir) / f"{file_hash}.parquet"

    @classmethod
    def pdf_to_text(cls, pdf_path: str, use_cache: bool = True) -> str:
        """Processa PDF com sistema de cache"""
        cache_path = cls.get_cache_path(pdf_path)
        
        if use_cache and cache_path.exists():
            return pd.read_parquet(cache_path)["text"].iloc[0]
        
        text = cls.extract_text(pdf_path)
        pd.DataFrame({"text": [text]}).to_parquet(cache_path)
        return text

    @classmethod
    def process_batch(cls, pdf_paths: list, max_workers: int = 4) -> list:
        """Processa múltiplos PDFs em paralelo com barra de progresso"""
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            return list(tqdm(executor.map(cls.pdf_to_text, pdf_paths), total=len(pdf_paths)))

# ==================== OPENAI INTEGRATION ====================
def get_openai_client() -> Optional[OpenAI]:
    """Retorna o cliente OpenAI configurado"""
    # Verifica session_state primeiro
    if hasattr(st.session_state, 'openai_api_key') and st.session_state.openai_api_key:
        return OpenAI(api_key=st.session_state.openai_api_key)
    
    # Depois verifica secrets.toml (para deploy)
    try:
        if 'openai_api_key' in st.secrets:
            return OpenAI(api_key=st.secrets.openai_api_key)
    except:
        pass
    
    return None

def generate_image_description(image_bytes: bytes) -> Optional[str]:
    """Gera descrição de imagem usando a API da OpenAI"""
    client = get_openai_client()
    if not client:
        st.error("API key da OpenAI não configurada")
        return None
    
    try:
        b64_image = base64.b64encode(image_bytes).decode('utf-8')
        response = client.chat.completions.create(
            model="gpt-4-turbo",
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": "Descreva esta imagem detalhadamente em português."},
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/png;base64,{b64_image}"
                            },
                        },
                    ],
                }
            ],
            max_tokens=500,
        )
        return response.choices[0].message.content
    except Exception as e:
        st.error(f"Erro ao gerar descrição: {str(e)}")
        return None

# ==================== DETECÇÃO DE TIPOS ====================
def detect_chunk_type(chunk: str) -> str:
    """Detecta o tipo de conteúdo do chunk"""
    if "```" in chunk:  # Se contém blocos de código
        return "código"
    elif any(char.isdigit() for char in chunk) and sum(c.isalpha() for c in chunk) / len(chunk) < 0.5:
        return "dados numéricos"
    elif len(chunk.split('\n')) > 5 and all(len(line) < 100 for line in chunk.split('\n')):
        return "lista"
    else:
        return "texto"

# ==================== FORMATAÇÃO ====================
def format_table(matrix: List[List]) -> str:
    """Formata uma matriz como tabela Markdown"""
    if not matrix or len(matrix) < 2:
        return ""
    
    # Cabeçalho
    markdown = "| " + " | ".join(str(cell) for cell in matrix[0]) + " |\n"
    markdown += "| " + " | ".join(["---"] * len(matrix[0])) + " |\n"
    
    # Linhas
    for row in matrix[1:]:
        markdown += "| " + " | ".join(str(cell) for cell in row) + " |\n"
    
    return markdown

# ==================== PROCESSADORES DE ARQUIVOS ====================
async def process_image_to_chunks(file_path: str) -> bool:
    """Processa imagens de forma assíncrona e extrai texto OCR."""
    try:
        img = Image.open(file_path)
        text = pytesseract.image_to_string(img)
        file_name = Path(file_path).name
        upload_time = datetime.now().strftime("%d/%m/%Y")

        async with aiosqlite.connect(DB_PATH) as db:
            if text.strip():
                chunks = textwrap.wrap(text, st.session_state.get('CHUNK_SIZE', CHUNK_SIZE))
                chunk_counter = 1
                for chunk in chunks:
                    chunk_type = detect_chunk_type(chunk)
                    markdown_chunk = f"**Arquivo**: {file_name}\n**Chunk**: {chunk_counter}\n\n{chunk.strip()}"
                    await db.execute('''
                        INSERT INTO chunks (file_name, page_number, chunk_number, content, chunk_content, upload_time, image_description, image_data)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                    ''', (file_name, None, chunk_counter, markdown_chunk, chunk_type, upload_time, None, None))
                    chunk_counter += 1
            await db.commit()
        return True
    except Exception as e:
        st.error(f"Erro ao processar imagem: {str(e)}")
        return False

async def process_pdf_to_chunks(file_path: str) -> None:
    """Processa arquivos PDF de forma assíncrona, extrai texto, tabelas e imagens."""
    try:
        doc = fitz.open(file_path)
        file_name = Path(file_path).name
        upload_time = datetime.now().strftime("%d/%m/%Y")

        async with aiosqlite.connect(DB_PATH) as db:
            tasks = []
            for page_index in range(len(doc)):
                page = doc.load_page(page_index)
                tasks.append(_process_pdf_page(db, page, file_name, upload_time, page_index))

            await asyncio.gather(*tasks)
            await db.commit()
    except Exception as e:
        st.error(f"Erro ao processar PDF: {str(e)}")

async def _process_pdf_page(db, page, file_name: str, upload_time: str, page_index: int) -> None:
    """Processa uma única página do PDF."""
    doc = page.parent  # Obtém o documento PDF da página atual
    text = page.get_text("text")
    tables = page.find_tables()
    images = page.get_images(full=True)
    chunk_counter = 1

    # Processa texto
    if text.strip():
        chunks = textwrap.wrap(text, st.session_state.get('CHUNK_SIZE', CHUNK_SIZE))
        for chunk in chunks:
            chunk_type = detect_chunk_type(chunk)
            markdown_chunk = f"**Arquivo**: {file_name}\n**Página**: {page_index + 1}\n**Chunk**: {chunk_counter}\n\n{chunk.strip()}"
            await db.execute('''
                INSERT INTO chunks (file_name, page_number, chunk_number, content, chunk_content, upload_time, image_description, image_data)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            ''', (file_name, page_index + 1, chunk_counter, markdown_chunk, chunk_type, upload_time, None, None))
            chunk_counter += 1

    # Processa tabelas
    if tables:
        for table in tables:
            matrix = table.extract()
            markdown_table = format_table(matrix)
            if markdown_table:
                markdown_chunk = f"**Arquivo**: {file_name}\n**Página**: {page_index + 1}\n**Chunk**: {chunk_counter}\n\n{markdown_table}"
                await db.execute('''
                    INSERT INTO chunks (file_name, page_number, chunk_number, content, chunk_content, upload_time, image_description, image_data)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                ''', (file_name, page_index + 1, chunk_counter, markdown_chunk, "tabela", upload_time, None, None))
                chunk_counter += 1

    # Processa imagens
    for img in images:
        await _process_pdf_image(db, doc, img, file_name, upload_time, page_index, chunk_counter)
        chunk_counter += 1

async def _process_pdf_image(db, doc, img, file_name: str, upload_time: str, page_index: int, chunk_counter: int) -> None:
    """Processa uma única imagem extraída do PDF."""
    try:
        xref = img[0]
        base_image = doc.extract_image(xref)
        image_bytes = base_image["image"]

        image_description = await _generate_image_description_async(image_bytes)
        b64_image = base64.b64encode(image_bytes).decode('ascii')
        image_meta = f"**Arquivo**: {file_name}\n**Página**: {page_index + 1}\n**Chunk**: {chunk_counter}\n**Descrição**: {image_description}"

        await db.execute('''
            INSERT INTO chunks (file_name, page_number, chunk_number, content, chunk_content, upload_time, image_description, image_data)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        ''', (file_name, page_index + 1, chunk_counter, image_meta, "imagem", upload_time, image_description, b64_image))
    except Exception as e:
        st.error(f"Erro ao processar imagem do PDF: {str(e)}")

async def _generate_image_description_async(image_bytes: bytes) -> str:
    """Gera descrição de imagem de forma assíncrona."""
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, generate_image_description, image_bytes)

async def process_docx_to_chunks(file_path: str) -> None:
    """Processa arquivos DOCX de forma assíncrona e extrai textos e tabelas."""
    try:
        doc = docx.Document(file_path)
        file_name = Path(file_path).name
        upload_time = datetime.now().strftime("%d/%m/%Y")

        async with aiosqlite.connect(DB_PATH) as db:
            chunk_counter = 1
            tasks = []

            # Processa parágrafos de texto
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    chunks = textwrap.wrap(text, st.session_state.get('CHUNK_SIZE', CHUNK_SIZE))
                    for chunk in chunks:
                        chunk_type = detect_chunk_type(chunk)
                        markdown_chunk = f"**Arquivo**: {file_name}\n**Chunk**: {chunk_counter}\n\n{chunk.strip()}"
                        tasks.append(db.execute('''
                            INSERT INTO chunks (file_name, page_number, chunk_number, content, chunk_content, upload_time, image_description, image_data)
                            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                        ''', (file_name, None, chunk_counter, markdown_chunk, chunk_type, upload_time, None, None)))
                        chunk_counter += 1

            # Processa tabelas do documento
            for table in doc.tables:
                matrix = []
                for row in table.rows:
                    matrix.append([cell.text.strip() for cell in row.cells])
                markdown_table = format_table(matrix)
                if markdown_table:
                    markdown_chunk = f"**Arquivo**: {file_name}\n**Chunk**: {chunk_counter}\n\n{markdown_table}"
                    tasks.append(db.execute('''
                        INSERT INTO chunks (file_name, page_number, chunk_number, content, chunk_content, upload_time, image_description, image_data)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                    ''', (file_name, None, chunk_counter, markdown_chunk, "tabela", upload_time, None, None)))
                    chunk_counter += 1

            if tasks:
                await asyncio.gather(*tasks)
                await db.commit()
    except Exception as e:
        st.error(f"Erro ao processar DOCX: {str(e)}")

async def process_xlsx_to_chunks(file_path: str) -> None:
    """Processa arquivos XLSX de forma assíncrona e extrai dados."""
    try:
        wb = openpyxl.load_workbook(file_path)
        file_name = Path(file_path).name
        upload_time = datetime.now().strftime("%d/%m/%Y")

        async with aiosqlite.connect(DB_PATH) as db:
            chunk_counter = 1
            tasks = []

            for sheet in wb.sheetnames:
                ws = wb[sheet]
                for row in ws.iter_rows(values_only=True):
                    row_data = ' | '.join([str(cell) if cell is not None else '' for cell in row])
                    if row_data.strip():
                        markdown_chunk = f"**Arquivo**: {file_name}\n**Planilha**: {sheet}\n**Chunk**: {chunk_counter}\n\n{row_data}"
                        tasks.append(db.execute('''
                            INSERT INTO chunks (file_name, page_number, chunk_number, content, chunk_content, upload_time, image_description, image_data)
                            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                        ''', (file_name, None, chunk_counter, markdown_chunk, "tabela", upload_time, None, None)))
                        chunk_counter += 1

            if tasks:
                await asyncio.gather(*tasks)
                await db.commit()
    except Exception as e:
        st.error(f"Erro ao processar XLSX: {str(e)}")

# ==================== EXPORTAÇÃO PARA MARKDOWN ====================
def export_chunk_to_md(chunk_id: int) -> Tuple[Optional[str], Optional[str]]:
    """Exporta um chunk específico para um arquivo Markdown"""
    try:
        with sqlite3.connect(DB_PATH) as conn:
            conn.row_factory = sqlite3.Row
            c = conn.cursor()
            c.execute("SELECT * FROM chunks WHERE id = ?", (chunk_id,))
            row = c.fetchone()
        
        if not row:
            return None, None
        
        md_content = _build_markdown_content(row)
        filename = f"chunk_{row['id']}_{row['file_name']}.md"
        
        return md_content, filename
    except Exception as e:
        st.error(f"Erro ao exportar chunk: {str(e)}")
        return None, None

def export_all_chunks_to_md(file_name: str) -> Tuple[Optional[str], Optional[str]]:
    """Exporta todos os chunks de um arquivo para um único arquivo Markdown"""
    try:
        with sqlite3.connect(DB_PATH) as conn:
            conn.row_factory = sqlite3.Row
            c = conn.cursor()
            c.execute("""
                SELECT * FROM chunks 
                WHERE file_name = ? 
                ORDER BY page_number, chunk_number
            """, (file_name,))
            rows = c.fetchall()
        
        if not rows:
            return None, None
        
        md_parts = [f"# {file_name}\n\n"]
        
        for row in rows:
            md_parts.append(f"---\n")
            md_parts.append(f"**ID:** {row['id']}\n\n")
            md_parts.append(f"**Nome do Arquivo:** {row['file_name']}\n\n")
            md_parts.append(f"**Página:** {row['page_number']}\n\n")
            md_parts.append(f"**Chunk:** {row['chunk_number']}\n\n")
            md_parts.append(f"**Tipo:** {row['chunk_content']}\n\n")
            md_parts.append(f"**Data de upload:** {row['upload_time']}\n\n")

            if row['chunk_content'] == "imagem":
                md_parts.append("## Descrição da Imagem\n\n")
                md_parts.append(f"{row['image_description']}\n\n")
                if row['image_data']:
                    md_parts.append(f"![Imagem](data:image/png;base64,{row['image_data']})\n\n")
            else:
                md_parts.append("## Conteúdo\n\n")
                md_parts.append(f"{row['content']}\n\n")
        
        return "".join(md_parts), f"documento_completo_{file_name}.md"
    except Exception as e:
        st.error(f"Erro ao exportar documento: {str(e)}")
        return None, None

def _build_markdown_content(row: sqlite3.Row) -> str:
    """Constrói o conteúdo Markdown para um chunk individual"""
    md_content = f"# {row['file_name']}\n\n"
    md_content += f"**ID:** {row['id']}\n"
    md_content += f"**Página:** {row['page_number']}\n"
    md_content += f"**Chunk:** {row['chunk_number']}\n"
    md_content += f"**Tipo:** {row['chunk_content']}\n"
    md_content += f"**Data de upload:** {row['upload_time']}\n\n"
    
    if row['chunk_content'] == "imagem":
        md_content += "## Descrição da Imagem\n"
        md_content += f"{row['image_description']}\n\n"
        md_content += f"![Imagem](data:image/png;base64,{row['image_data']})"
    else:
        md_content += row['content']
    
    return md_content
